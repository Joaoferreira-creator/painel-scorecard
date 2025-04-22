# Painel ScoreCard - Aplicativo Profissional
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import datetime
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== Configurações Iniciais =====
st.set_page_config(page_title="Painel ScoreCard", layout="wide")

# Carregar o Logo
logo = Image.open("REMOTA.png")

# Cabeçalho Colorido com Logo
col1, col2 = st.columns([1, 5])
with col1:
    st.image(logo, width=120)
with col2:
    st.markdown(
        """
        <div style="background-color: #0077b6; padding: 15px; border-radius: 10px;">
            <h1 style="color: white;">Painel ScoreCard</h1>
            <p style="color: white;">
            Avaliação do suporte institucional, gerencial e tecnológico em ambientes remotos/híbridos.
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

st.markdown("---")

# Sidebar
st.sidebar.header("Menu Principal")
st.sidebar.image(logo, width=180)

# Upload do arquivo
arquivo = st.sidebar.file_uploader("Selecione seu arquivo (.xlsx)", type=["xlsx"])
opcao = st.sidebar.radio("Escolha uma seção:", ["Classificação Geral", "Gráficos Individuais", "Tabelas de Cruzamento", "Gerar Relatório"])

# ===== Processamento de Dados =====
if arquivo:
    df = pd.read_excel(arquivo)

    # Dimensões e Itens
    dimensoes = {
        'Política Institucional': ['pi1', 'pi2', 'pi3', 'pi4', 'pi5', 'pi6', 'pi7'],
        'Gestão de Desempenho': ['gd1', 'gd2', 'gd3', 'gd4', 'gd5', 'gd6', 'gd7', 'gd8'],
        'Suporte Gestor Projeto': ['sg1', 'sg2', 'sg3', 'sg4', 'sg5', 'sg6', 'sg7', 'sg8', 'sg9'],
        'Suporte Saúde Mental/Física': ['sm1', 'sm2', 'sm3', 'sm4', 'sm5', 'sm6', 'sm7', 'sm8', 'sm9', 'sm10', 'sm11', 'sm12', 'sm13', 'sm14', 'sm15'],
        'Ferramentas Tecnológicas': ['ft1', 'ft2', 'ft3', 'ft4', 'ft5', 'ft6'],
        'Tomada de Decisão': ['td1', 'td2', 'td3', 'td4', 'td5', 'td6', 'td7', 'td8']
    }

    variaveis_perfil = ['Faixa_idade', 'Faixa_renda', 'Sexo', 'cargo']
    cores = {'Bronze': 'red', 'Prata': 'orange', 'Ouro': 'green'}

    # Soma dos Itens de Cada Dimensão
    for dim, itens in dimensoes.items():
        df[f'Soma_{dim}'] = df[itens].sum(axis=1)

    # Função para Classificar
    def classificar_selo(valor, n_itens):
        minimo = n_itens * 1
        maximo = n_itens * 5
        intervalo = maximo - minimo
        q1 = minimo + 0.25 * intervalo
        q3 = minimo + 0.75 * intervalo
        if valor <= q1:
            return 'Bronze'
        elif valor <= q3:
            return 'Prata'
        else:
            return 'Ouro'

    # Calcular as Medianas
    medianas_dimensao = {dim: df[f'Soma_{dim}'].median() for dim in dimensoes.keys()}
    selo_dimensao = {dim: classificar_selo(medianas_dimensao[dim], len(itens)) for dim, itens in dimensoes.items()}

    for dim, itens in dimensoes.items():
        for item in itens:
            df[f'Classificacao_{item}'] = df[item].apply(lambda x: 'Bronze' if x <= 2 else ('Prata' if x <= 4 else 'Ouro'))

    # Classificação Geral
    minimos_teoricos = [len(itens)*1 for itens in dimensoes.values()]
    maximos_teoricos = [len(itens)*5 for itens in dimensoes.values()]
    mediana_min = pd.Series(minimos_teoricos).median()
    mediana_max = pd.Series(maximos_teoricos).median()
    intervalo_geral = mediana_max - mediana_min
    q1_geral = mediana_min + 0.25 * intervalo_geral
    q3_geral = mediana_min + 0.75 * intervalo_geral

    mediana_geral = pd.Series(list(medianas_dimensao.values())).median()

    if mediana_geral <= q1_geral:
        selo_geral = 'Bronze'
    elif mediana_geral <= q3_geral:
        selo_geral = 'Prata'
    else:
        selo_geral = 'Ouro'

    # ===== Telas do App =====
    if opcao == "Classificação Geral":
        st.header("Classificação de Selos por Dimensão + Selo Geral")

        dimensoes_list = list(medianas_dimensao.keys())
        medianas_list = list(medianas_dimensao.values())
        cores_barras = [cores[selo_dimensao[dim]] for dim in dimensoes_list]

        dimensoes_list.append('Selo Geral')
        medianas_list.append(mediana_geral)
        cores_barras.append(cores[selo_geral])

        fig, ax = plt.subplots(figsize=(12,7))
        barras = ax.bar(dimensoes_list, medianas_list, color=cores_barras)
        ax.set_title('Classificação de Selos por Dimensão + Selo Geral')
        ax.set_ylabel('Mediana da Soma das Respostas')
        ax.set_xlabel('Dimensão')
        plt.xticks(rotation=45, ha='right')
        plt.grid(axis='y')

        itens_dimensao = {k: len(v) for k, v in dimensoes.items()}

        for barra, dim, valor in zip(barras, dimensoes_list, medianas_list):
            altura = barra.get_height()
            if dim != 'Selo Geral':
                n_itens = itens_dimensao[dim]
                minimo = n_itens * 1
                maximo = n_itens * 5
            else:
                minimo = 8
                maximo = 40
            intervalo = maximo - minimo
            q1 = minimo + 0.25 * intervalo
            q3 = minimo + 0.75 * intervalo

            if valor <= q1:
                faixa = "≤25%"
            elif valor <= q3:
                faixa = "25%-75%"
            else:
                faixa = ">75%"

            ax.annotate(f'{valor:.1f} ({faixa})', xy=(barra.get_x() + barra.get_width()/2, altura),
                        xytext=(0, 6), textcoords='offset points', ha='center', va='bottom')

        legenda_bronze = mpatches.Patch(color='red', label='Bronze')
        legenda_prata = mpatches.Patch(color='orange', label='Prata')
        legenda_ouro = mpatches.Patch(color='green', label='Ouro')
        plt.legend(handles=[legenda_bronze, legenda_prata, legenda_ouro],
                   title="Legenda", loc='center left', bbox_to_anchor=(1.02, 0.5), borderaxespad=0)

        st.pyplot(fig)

    elif opcao == "Gráficos Individuais":
        st.header("Distribuição dos Respondentes por Dimensão e por Item")

        dimensao_escolhida = st.selectbox("Selecione uma Dimensão:", list(dimensoes.keys()))
        itens_disponiveis = dimensoes[dimensao_escolhida]
        itens_escolhidos = st.multiselect("Selecione os Itens:", itens_disponiveis, default=itens_disponiveis)

        if itens_escolhidos:
            for item in itens_escolhidos:
                contagem = df[f'Classificacao_{item}'].value_counts(normalize=True) * 100
                contagem = contagem.reindex(['Bronze', 'Prata', 'Ouro']).fillna(0)

                fig, ax = plt.subplots()
                ax.bar(contagem.index, contagem.values, color=[cores[c] for c in contagem.index])
                ax.set_title(f'Distribuição dos Respondentes - {item}')
                ax.set_ylabel('Percentual (%)')
                ax.set_ylim(0, 100)

                for i, v in enumerate(contagem.values):
                    ax.text(i, v + 2, f'{v:.1f}%', ha='center', va='bottom')

                st.pyplot(fig)

    elif opcao == "Tabelas de Cruzamento":
        st.header("Tabelas de Cruzamento Perfil x Classificação")

        df_selos = pd.DataFrame()
        for dim, itens in dimensoes.items():
            df_selos[dim] = df[f'Soma_{dim}'].apply(lambda x: classificar_selo(x, len(itens)))

        for var in variaveis_perfil:
            df_selos[var] = df[var]

        for dimensao in dimensoes.keys():
            for variavel in variaveis_perfil:
                st.subheader(f"{dimensao} x {variavel}")
                tabela = pd.crosstab(df_selos[variavel], df_selos[dimensao], normalize='index') * 100
                tabela = tabela.reindex(columns=['Bronze', 'Prata', 'Ouro']).fillna(0)
                st.dataframe(tabela.round(1))

    elif opcao == "Gerar Relatório":
        st.header("Gerar Relatório Técnico ScoreCard (.docx)")

        data_hoje = datetime.datetime.now().strftime('%Y-%m-%d')
        nome_arquivo = f"Relatorio_Geral_ScoreCard_{data_hoje}.docx"

        if st.button("📄 Gerar Relatório Word"):
            df_selos = pd.DataFrame()
            for dim, itens in dimensoes.items():
                df_selos[dim] = df[f'Soma_{dim}'].apply(lambda x: classificar_selo(x, len(itens)))
            for var in variaveis_perfil:
                df_selos[var] = df[var]

            doc = Document()

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("Relatório Geral - Escala Remota ScoreCard\n")
            run.bold = True
            run.font.size = Pt(24)

            run = p.add_run(f"Data de Geração: {data_hoje}")
            run.font.size = Pt(12)

            doc.add_page_break()

            doc.add_heading('1. Introdução', level=1)
            doc.add_paragraph("A Escala Remota ScoreCard avalia a qualidade do suporte institucional, gerencial e tecnológico em ambientes de trabalho remoto e híbrido.")

            doc.add_heading('2. Classificação Geral', level=1)
            for dim, med in medianas_dimensao.items():
                doc.add_paragraph(f"- {dim}: Mediana = {med:.1f} / Selo = {selo_dimensao[dim]}")
            doc.add_paragraph(f"**Classificação Geral:** {selo_geral}")

            doc.add_heading('3. Tabelas de Cruzamento Perfil x Selo', level=1)
            for dimensao in dimensoes.keys():
                for variavel in variaveis_perfil:
                    tabela = pd.crosstab(df_selos[variavel], df_selos[dimensao], normalize='index') * 100
                    tabela = tabela.reindex(columns=['Bronze', 'Prata', 'Ouro']).fillna(0)

                    doc.add_heading(f"{dimensao} x {variavel}", level=2)
                    table = doc.add_table(rows=1, cols=len(tabela.columns)+1)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = variavel
                    for idx, col in enumerate(tabela.columns):
                        hdr_cells[idx+1].text = col

                    for idx, index in enumerate(tabela.index):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(index)
                        for jdx, col in enumerate(tabela.columns):
                            row_cells[jdx+1].text = f"{tabela.loc[index, col]:.1f}%"

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="📥 Baixar Relatório Word (.docx)",
                data=buffer,
                file_name=nome_arquivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
